# Gera docs/guia-equipe.docx a partir de docs/guia-equipe.html.
# Converte o guia (HTML estilizado) num documento Word limpo, com a cor da
# clínica nos títulos. Rode: python docs/build_guia_equipe_docx.py
import os
import re
from html.parser import HTMLParser

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, 'guia-equipe.html')
OUT = os.path.join(HERE, 'guia-equipe.docx')

TEAL = RGBColor(0x0F, 0x76, 0x6E)
SLATE = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x64, 0x74, 0x8B)

# Blocos que viram parágrafos, na ordem em que aparecem no HTML.
BLOCK_TAGS = {'h1', 'h2', 'h3', 'p', 'li', 'div', 'footer'}


class Bloco:
    """Um bloco de texto extraído do HTML: tipo + trechos (texto, negrito)."""

    def __init__(self, kind):
        self.kind = kind
        self.runs = []          # [(texto, bold)]

    def texto(self):
        return ''.join(t for t, _ in self.runs).strip()


class GuiaParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocos = []
        self.atual = None
        self.bold = 0
        self.skip = 0           # dentro de <style>/<script>/<nav>
        self.pilha = []         # (tag, classe)
        self.em_steps = False
        self.em_ol = False      # dentro de qualquer <ol> (lista numerada)

    # --- helpers ----------------------------------------------------------
    def _classe(self, attrs):
        return dict(attrs).get('class', '')

    def _kind(self, tag, classe):
        if tag in ('h1', 'h2', 'h3'):
            return tag if tag != 'h2' else ('group' if 'group-title' in classe else 'h2')
        if tag == 'li':
            return 'step' if (self.em_steps or self.em_ol) else 'bullet'
        if tag == 'footer':
            return 'footer'
        if tag == 'div':
            if 'where' in classe:
                return 'where'
            if 'callout' in classe:
                return 'callout'
            return None          # divs estruturais não geram parágrafo
        # <p>
        for c in ('eyebrow', 'group-eyebrow', 'lede', 'meta', 'purpose', 'block-label', 'note'):
            if c in classe:
                return c
        return 'p'

    # --- eventos ----------------------------------------------------------
    def handle_starttag(self, tag, attrs):
        classe = self._classe(attrs)
        if tag in ('style', 'script') or (tag == 'nav'):
            self.skip += 1
            return
        if self.skip:
            return
        self.pilha.append((tag, classe))
        if tag == 'ol':
            self.em_ol = True
            if 'steps' in classe:
                self.em_steps = True
        if tag in ('b', 'strong'):
            self.bold += 1
            return
        if tag in BLOCK_TAGS:
            kind = self._kind(tag, classe)
            if kind:
                self.atual = Bloco(kind)
                self.blocos.append(self.atual)

    def handle_endtag(self, tag):
        if tag in ('style', 'script', 'nav'):
            self.skip = max(0, self.skip - 1)
            return
        if self.skip:
            return
        if tag == 'ol':
            self.em_steps = False
            self.em_ol = False
        if tag in ('b', 'strong'):
            self.bold = max(0, self.bold - 1)
            return
        if tag in BLOCK_TAGS:
            self.atual = None
        if self.pilha:
            self.pilha.pop()

    def handle_data(self, data):
        if self.skip or self.atual is None:
            return
        texto = re.sub(r'\s+', ' ', data)
        if texto.strip() == '':
            # preserva espaço entre runs (ex.: "<b>Onde fica:</b> menu")
            if self.atual.runs and not self.atual.runs[-1][0].endswith(' '):
                self.atual.runs.append((' ', False))
            return
        self.atual.runs.append((texto, self.bold > 0))


def sombrear(par, hexcor):
    """Fundo suave no parágrafo (usado nos callouts)."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcor)
    par._p.get_or_add_pPr().append(shd)


def escrever(doc, bloco):
    kind, texto = bloco.kind, bloco.texto()
    if not texto:
        return

    def par(space_before=0, space_after=4, esquerda=0, style=None):
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if esquerda:
            p.paragraph_format.left_indent = Pt(esquerda)
        return p

    def runs_em(p, size=10.5, color=SLATE, italic=False, caps=False):
        for t, bold in bloco.runs:
            if not t:
                continue
            r = p.add_run(t.upper() if caps else t)
            r.bold = bold
            r.italic = italic
            r.font.size = Pt(size)
            r.font.color.rgb = color
        return p

    # Títulos usam os estilos nativos (habilita o painel de navegação do Word),
    # com tamanho/cor da identidade da clínica.
    if kind == 'h1':
        p = par(0, 2, style='Title')
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(24)
        r.font.color.rgb = SLATE
    elif kind == 'group':
        p = par(0, 8, style='Heading 1')
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = TEAL
    elif kind == 'h2':
        p = par(10, 6, style='Heading 1')
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = TEAL
    elif kind == 'h3':
        p = par(12, 3, style='Heading 2')
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = SLATE
    elif kind == 'group-eyebrow':
        # Início de área: quebra de página antes do rótulo, para o bloco ficar junto.
        doc.add_page_break()
        runs_em(par(0, 1), size=8, color=TEAL, caps=True)
    elif kind == 'eyebrow':
        runs_em(par(8, 1), size=8, color=TEAL, caps=True)
    elif kind == 'lede':
        runs_em(par(4, 6), size=11, color=GRAY)
    elif kind in ('meta', 'note'):
        runs_em(par(2, 6), size=8.5, color=GRAY, italic=True)
    elif kind == 'purpose':
        runs_em(par(0, 4), size=10.5, color=GRAY, italic=True)
    elif kind == 'where':
        runs_em(par(0, 4), size=9.5, color=GRAY)
    elif kind == 'block-label':
        p = par(6, 2)
        r = p.add_run(texto.upper())
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = TEAL
    elif kind == 'callout':
        p = par(4, 6, esquerda=8)
        runs_em(p, size=9.5, color=SLATE)
        sombrear(p, 'F1F5F9')
    elif kind in ('step', 'bullet'):
        p = doc.add_paragraph(style='List Number' if kind == 'step' else 'List Bullet')
        p.paragraph_format.space_after = Pt(2)
        runs_em(p, size=10.5)
    elif kind == 'footer':
        p = par(14, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runs_em(p, size=8.5, color=GRAY, italic=True)
    else:  # parágrafo comum
        runs_em(par(0, 5))


def main():
    with open(SRC, encoding='utf-8') as f:
        html = f.read()

    parser = GuiaParser()
    parser.feed(html)

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Pt(48)
        s.left_margin = s.right_margin = Pt(54)

    for bloco in parser.blocos:
        escrever(doc, bloco)

    doc.save(OUT)
    print(f'DOCX gerado: {OUT} - {os.path.getsize(OUT)} bytes ({len(parser.blocos)} blocos)')


if __name__ == '__main__':
    main()
